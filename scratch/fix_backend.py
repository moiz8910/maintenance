import os

file_path = "backend/main.py"
with open(file_path, "r", encoding="utf-8") as f:
    lines = f.readlines()

cut_start = -1
for i, line in enumerate(lines):
    if '"requester_department": "Maintenance Division",' in line:
        cut_start = i + 1
        break

if cut_start == -1:
    print("Could not find cut start")
    exit(1)

cut_end = -1
for i in range(cut_start, len(lines)):
    if '@app.get("/api/inventory/pending-prs")' in lines[i]:
        cut_end = i
        break

if cut_end == -1:
    print("Could not find cut end")
    exit(1)

restored_code = """            "justification": f"Urgent replacement of {material} to restore {asset_name} operational integrity.",
            "technical_specifications": "Standard industrial grade specifications for aluminum smelter environments.",
            "vendor_recommendations": ["Global Industrial Supplies", "Vedanta Approved Local Vendors"],
            "estimated_budget": "₹85,000.00 (Estimated)",
            "delivery_urgency": "IMMEDIATE",
            "inspection_requirements": "Visual inspection for transit damage and verification of mill test certificates.",
            "approval_workflow": "Maintenance Lead -> Procurement Manager -> Finance Head"
        }

@app.post("/api/purchase-requisition/download-docx")
async def download_pr_docx(data: dict):
    \"\"\"Convert provided PR JSON data into a Word document.\"\"\"
    import io
    from docx import Document as DocxDocument
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pr_no = data.get("pr_number", "PR-NEW")
    mat = data.get("matName", "Unknown Material")

    doc = DocxDocument()
    
    # Title
    title = doc.add_heading("VEDANTA JHARSUGUDA — PURCHASE REQUISITION", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph(f"PR NUMBER: {pr_no}  |  MATERIAL: {mat}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Department: {data.get('requester_department','')}   |   Urgency: {data.get('delivery_urgency','')}   |   Budget: {data.get('estimated_budget','')}")
    doc.add_paragraph()

    def add_section(heading, content):
        doc.add_heading(heading, level=2)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(f"• {item}", style="List Bullet")
        else:
            doc.add_paragraph(str(content))

    sections = [
        ("Technical Justification", "justification"),
        ("Technical Specifications", "technical_specifications"),
        ("Vendor Recommendations", "vendor_recommendations"),
        ("QA/QC & Inspection Requirements", "inspection_requirements"),
        ("Approval Workflow", "approval_workflow")
    ]

    for title_text, key in sections:
        add_section(title_text, data.get(key, ""))

    doc.add_paragraph()
    doc.add_heading("Authorization Signatures", level=2)
    for role in ["Maintenance Lead", "Procurement Manager", "Finance Director", "Smelter Operations Head"]:
        doc.add_paragraph(f"{role}: ___________________________    Date: ____________")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\\"PurchaseRequisition_{pr_no}.docx\\""}
    )

class UpdateLeadTimeRequest(BaseModel):
    material_id: str
    lead_time: int

@app.post("/api/work-order/{wo_id}/update-material-lead-time")
async def update_material_lead_time(wo_id: str, payload: UpdateLeadTimeRequest):
    conn = sqlite3.connect(settings.db_path)
    cursor = conn.cursor()
    cursor.execute(\"\"\"
        UPDATE task_material_linkage 
        SET lead_time = ? 
        WHERE material_used = ? AND work_order_task_item IN (
            SELECT id FROM work_order_task_item WHERE work_order = ?
        )
    \"\"\", (payload.lead_time, payload.material_id, wo_id))
    conn.commit()
    conn.close()
    
    # Reschedule the work order
    new_date = reschedule_work_order(wo_id)
    return {"status": "success", "new_date": new_date}

@app.get("/api/inventory/summary")
async def get_inventory_summary():
    conn = sqlite3.connect(settings.db_path)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # 1. Total Value in INR
    cursor.execute(\"\"\"
        SELECT SUM(oh.stock_available_on_hand * mp.price_per_unit) as total_value
        FROM on_hand_inventory oh
        JOIN material_price mp ON oh.material = mp.material
    \"\"\")
    total_value = cursor.fetchone()['total_value'] or 0
    
    # 2. Pending PR Count
    cursor.execute("SELECT COUNT(*) as count FROM purchase_requisition WHERE LOWER(status) = 'pending'")
    pending_pr_count = cursor.fetchone()['count'] or 0
    
    # 3. Critical Spares Out of Stock (Spares with 0 stock)
    cursor.execute(\"\"\"
        SELECT COUNT(DISTINCT mm.id) as count
        FROM material_master mm
        JOIN on_hand_inventory oh ON mm.id = oh.material
        WHERE mm.material_type = 'Spares' AND oh.stock_available_on_hand <= 0
    \"\"\")
    critical_oos_count = cursor.fetchone()['count'] or 0
    
    # 4. Obsolescence Count
    cursor.execute(\"\"\"
        SELECT oh.receipt_date, mm.shelf_life
        FROM on_hand_inventory oh
        JOIN material_master mm ON oh.material = mm.id
        WHERE mm.shelf_life IS NOT NULL AND oh.receipt_date IS NOT NULL
    \"\"\")
    rows = cursor.fetchall()
    obsolescence_count = 0
    from datetime import datetime, timedelta
    today = datetime.now()
    threshold = today + timedelta(days=30)
    
    for row in rows:
        try:
            r_date = datetime.strptime(row['receipt_date'], "%d-%m-%y")
            expiry_date = r_date + timedelta(days=row['shelf_life'])
            if expiry_date <= threshold:
                obsolescence_count += 1
        except:
            continue

    conn.close()
    return {
        "total_value_inr": total_value,
        "pending_pr_count": pending_pr_count,
        "critical_oos_count": critical_oos_count,
        "obsolescence_count": obsolescence_count
    }

"""

new_lines = lines[:cut_start] + [restored_code] + lines[cut_end:]

with open(file_path, "w", encoding="utf-8") as f:
    f.writelines(new_lines)

print("File fixed successfully")
