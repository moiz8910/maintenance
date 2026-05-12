import os
import io
import json
import asyncio
import sqlite3
from typing import List, Optional
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from pydantic_settings import BaseSettings
from google import genai
from langgraph.graph import StateGraph, END
from datetime import datetime, timedelta
from dotenv import load_dotenv

# Import our custom engines
from services.kpi_engine import get_all_kpis
from services.drilldown_engine import get_drilldown_data, safe_query
from services.agent_manager import run_maintenance_assistant, run_agent_workflow, generate_with_retry
from services.scheduling_engine import reschedule_work_order
from lookup_router import router as lookup_router

load_dotenv()

class Settings(BaseSettings):
    google_api_key: str = os.getenv("GOOGLE_API_KEY", "")
    openai_api_key: str = os.getenv("OPENAI_API_KEY", "")
    model_provider: str = os.getenv("MODEL_PROVIDER", "google")
    db_path: str = os.path.join(os.path.dirname(os.path.abspath(__file__)), "maintenance.db")
    model_config = {"env_file": ".env", "extra": "allow"}

settings = Settings()

app = FastAPI(title="AI-Powered Maintenance Intelligence Platform")
app.include_router(lookup_router)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000", 
        "http://127.0.0.1:3000", 
        "http://0.0.0.0:3000",
        "http://localhost:8000",
        "http://127.0.0.1:8000"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Models
class KPI(BaseModel):
    name: str
    value: str
    status: str
    subtext: str

class ChatRequest(BaseModel):
    message: str
    context_data: Optional[dict] = None

# Log helper
def log_stage(stage: int, message: str):
    print(f"[Stage {stage}] {message}")

# AI Client
client = genai.Client(api_key=settings.google_api_key)

# ─── ENDPOINTS ───

@app.get("/api/kpis", response_model=List[KPI])
async def get_kpis():
    log_stage(4, "Computing KPIs from real data...")
    return get_all_kpis()

@app.get("/api/drilldown/{kpi_id}")
async def get_drilldown(kpi_id: str):
    # kpi_id comes in slug-format from frontend (e.g. "work-order")
    return get_drilldown_data(kpi_id)

@app.get("/api/assets")
async def get_assets():
    return safe_query("""
        SELECT
            a.id,
            a.name,
            at.type          AS asset_type,
            a.location,
            a.criticality,
            a.throughput_rate,
            a.throughput_rate_uom,
            a.mean_time_to_repairmttr_value   AS mttr,
            a.mean_time_to_repairmttr_uom     AS mttr_uom,
            a.mean_time_between_failuresmtbf  AS mtbf,
            a.mean_time_between_failuresmtbf_uom AS mtbf_uom,
            a.unplanned_downtime,
            a.unplanned_downtime_uom,
            a.sop_number,
            a.sop_description,
            p.name           AS parent_name,
            p.id             AS parent_id,
            p.location       AS parent_location
        FROM asset a
        LEFT JOIN asset_type at ON a.type = at.id
        LEFT JOIN asset p ON a.parent_asset = p.id
        ORDER BY a.id
    """)

@app.get("/api/work-orders")
async def get_work_orders(status: Optional[str] = None):
    query = """
        SELECT w.id, w.repair_description as description, w.work_order_class as class, w.work_order_status as status,
               CASE WHEN EXISTS (SELECT 1 FROM work_order_task_item t WHERE t.work_order = w.id) THEN 1 ELSE 0 END as has_task,
               w.work_order_open_day as date,
               (SELECT t.work_order_task_item_open_day FROM work_order_task_item t WHERE t.work_order = w.id LIMIT 1) as schedule_date,
               CASE WHEN EXISTS (
                   SELECT 1 
                   FROM task_material_linkage m
                   JOIN work_order_task_item ti ON m.work_order_task_item = ti.id
                   LEFT JOIN on_hand_inventory inv ON inv.material = m.material_used
                   WHERE ti.work_order = w.id
                   GROUP BY m.id
                   HAVING m.quantity_used > COALESCE(SUM(inv.stock_available_on_hand), 0)
               ) THEN 1 ELSE 0 END as pr_needed
        FROM work_order w
    """
    if status:
        query += f" WHERE LOWER(w.work_order_status) = '{status.lower()}'"
    query += " ORDER BY pr_needed DESC, has_task DESC, w.work_order_class ASC"
    return safe_query(query)

@app.get("/api/diagnostic/work-orders")
async def get_diagnostic_work_orders():
    """Fetch all work orders currently in the 'Diagnostic' state."""
    log_stage(15, "Fetching Diagnostic Queue...")
    query = """
        SELECT w.id, w.repair_description as description, w.work_order_class as class, 
               w.work_order_status as status, w.repair_type as type,
               w.asset_id as asset, a.name as asset_name, w.work_order_open_day as date
        FROM work_order w
        LEFT JOIN asset a ON w.asset_id = a.id
        WHERE LOWER(w.work_order_status) = 'diagnostic'
        ORDER BY w.work_order_open_day DESC
    """
    return safe_query(query)

@app.get("/api/drilldown/safety-incidents")
async def get_drilldown_incidents(filter: Optional[str] = None):
    query = "SELECT * FROM incident_events"
    if filter and filter != 'All':
        query += f" WHERE incident_type = '{filter}'"
    return safe_query(query)

@app.get("/api/drilldown/safety-compliance")
async def get_drilldown_compliance(filter: Optional[str] = None):
    query = "SELECT id, work_order, asset, type, status FROM work_permit"
    if filter and filter != 'All':
        query += f" WHERE type = '{filter}'"
    return safe_query(query)

@app.get("/api/drilldown/manpower-utilization")
async def get_drilldown_technicians(filter: Optional[str] = None):
    query = "SELECT name, role_designation as role, discipline_trade as discipline, standard_hourly_rate as rate FROM technician_engineer"
    if filter and filter != 'All':
        query += f" WHERE role_designation = '{filter}' OR discipline_trade = '{filter}'"
    return safe_query(query)

@app.get("/api/drilldown/purchase-requisition")
async def get_drilldown_prs(filter: Optional[str] = None):
    query = "SELECT id, material_name as material, quantity, status FROM purchase_requisition"
    if filter and filter != 'All':
        query += f" WHERE LOWER(status) = '{filter.lower()}'"
    return safe_query(query)

@app.post("/api/work-order/{wo_id}/diagnose")
async def run_diagnosis(wo_id: str):
    """Use AI to analyze a breakdown and suggest tasks."""
    # 1. Fetch WO and History
    wo_details = safe_query("SELECT repair_description, asset_id as asset, repair_type FROM work_order WHERE id = ?", (wo_id,))
    if not wo_details:
        raise HTTPException(status_code=404, detail="Work Order not found")
    
    asset_id = wo_details[0]['asset']
    history = safe_query("""
        SELECT repair_description, work_order_status 
        FROM work_order 
        WHERE asset = ? AND id != ? AND work_order_status = 'Closed'
        LIMIT 5
    """, (asset_id, wo_id))
    
    # 2. Construct Prompt
    prompt = f"""
    Act as a Senior Maintenance Diagnostic Expert for an Aluminum Smelter.
    
    Current Problem: {wo_details[0]['repair_description']}
    Repair Type: {wo_details[0]['repair_type']}
    Asset ID: {asset_id}
    
    Past History for this Asset:
    {json.dumps(history)}
    
    CRITICAL: 
    - Provide a definitive, technical Root Cause Analysis.
    - DO NOT use phrases like "in the absence of history", "based on limited data", or "it's difficult to say".
    - Act as if you have seen this problem a thousand times and provide your most probable engineering judgment.
    
    Return your response EXACTLY in this JSON format:
    {{
      "probable_cause": "description here",
      "suggested_tasks": ["Task 1", "Task 2", "Task 3"]
    }}
    """
    from services.agent_manager import generate_with_retry
    result_str = generate_with_retry(prompt)
    print(f"[Diagnosis API] Raw Response: {result_str}")
    
    # Simple JSON extraction
    try:
        import re
        json_match = re.search(r'\{.*\}', result_str, re.DOTALL)
        if json_match:
            result = json.loads(json_match.group())
        else:
            result = json.loads(result_str)
        print(f"[Diagnosis API] Parsed Result: {result}")
        # Add original description to result
        result["reported_issue"] = wo_details[0].get("repair_description")
        return result
    except Exception as e:
        print(f"[Diagnosis API] Parse Error: {e}")
        return {
            "probable_cause": "AI failed to parse response.", 
            "suggested_tasks": ["Manual inspection required"],
            "reported_issue": wo_details[0].get("repair_description")
        }

@app.post("/api/work-order/{wo_id}/approve-diagnosis")
async def approve_diagnosis(wo_id: str, payload: dict):
    """Save approved tasks and move WO to Pending."""
    tasks = payload.get("tasks", [])
    asset_id = payload.get("asset_id")
    
    print(f"[Approval API] Starting approval for {wo_id} (Asset: {asset_id})")
    
    import sqlite3
    import os
    db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "maintenance.db")
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    try:
        # 1. Add Task Items
        for idx, task_desc in enumerate(tasks):
            ti_id = f"WOT-{wo_id[3:]}-{idx+1}"
            cursor.execute("""
                INSERT INTO work_order_task_item 
                (id, work_order, asset, task, work_order_task_item_open_day)
                VALUES (?, ?, ?, ?, ?)
            """, (ti_id, wo_id, asset_id, task_desc, datetime.now().strftime("%d-%m-%y")))
            
        # 2. Update WO Status
        cursor.execute("UPDATE work_order SET work_order_status = 'Pending' WHERE id = ?", (wo_id,))
        conn.commit()
        
        # 3. Trigger Auto-Pilot to populate the plan (Manpower, Materials, etc.)
        from services.agent_manager import run_agent_workflow
        print(f"[Approval API] Triggering Auto-Pilot for {wo_id}...")
        try:
            # We pass a message that specifically asks to plan this WO
            run_agent_workflow("maintenance_auto_pilot", f"Please generate a complete execution plan for {wo_id} based on its newly diagnosed tasks.")
        except Exception as ae:
            print(f"[Approval API] Auto-Pilot Warning: {ae}")
            # We don't fail the whole request if auto-pilot has a hiccup, 
            # but it usually works fine.

        print(f"[Approval API] SUCCESS: {wo_id} moved to Pending and planned.")
        return {"status": "success", "tasks_added": len(tasks)}
    except Exception as e:
        print(f"[Approval API] ERROR: {e}")
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@app.post("/api/work-order/{wo_id}/execution-advice")
async def get_execution_advice(wo_id: str):
    log_stage(25, f"Generating OpenAI Execution Advice for WO {wo_id}")
    
    # 1. Fetch context (needed for asset_id and AI prompt)
    wos = safe_query("""
        SELECT w.id, w.repair_description, w.repair_type, w.work_order_class,
               a.id as asset_id, a.name as asset_name, at.type as asset_type, a.location, a.criticality,
               a.sop_description
        FROM work_order w
        LEFT JOIN work_order_task_item woti ON woti.work_order = w.id
        LEFT JOIN asset a ON woti.asset = a.id
        LEFT JOIN asset_type at ON a.type = at.id
        WHERE w.id = ?
        LIMIT 1
    """, (wo_id,))
    
    if not wos:
        raise HTTPException(status_code=404, detail="Work order not found")
    wo = wos[0]

    # 2. Check Cache
    cached = safe_query("SELECT advice FROM execution_advice_cache WHERE wo_id = ?", (wo_id,))
    if cached:
        return {"advice": cached[0]['advice'], "asset_id": wo['asset_id'], "cached": True}

    # Fetch tasks for more context
    tasks = safe_query("SELECT task FROM work_order_task_item WHERE work_order = ?", (wo_id,))
    tasks_str = "\n".join([f"- {t['task']}" for t in tasks])

    prompt = f"""
    Act as a Senior Maintenance Engineer at Vedanta Jharsuguda Aluminum Smelter.
    Provide a detailed Execution Strategy and Advice for the following Work Order.
    
    Context:
    Work Order ID: {wo['id']}
    Asset ID: {wo['asset_id']}
    Asset Name: {wo['asset_name']}
    Asset Type: {wo['asset_type']}
    Description: {wo['repair_description']}
    Type: {wo['repair_type']}
    Class: {wo['work_order_class']}
    Location: {wo['location']}
    Asset Criticality: {wo['criticality']}
    Existing SOP Info: {wo['sop_description']}
    
    Planned Tasks:
    {tasks_str}
    
    Structure your response with:
    # Execution Strategy: {wo['id']}
    **Asset ID: {wo['asset_id']}**
    
    ## Technical Execution Steps
    - Provide a chronological, technical step-by-step guide.
    - Focus on precision (e.g., torque values, alignment tolerances).
    
    ## Safety & Compliance
    - List specific LOTO (Lockout/Tagout) requirements.
    - PPE requirements (Industrial standard).
    - Specific Life-Saving Rules applicable here.
    
    ## Tooling & Equipment
    - List all specialized tools (e.g., laser alignment kit, hydraulic pullers).
    - Consumables needed.
    
    ## Quality Control & OEM Standards
    - Technical specs to verify post-repair.
    
    ## Resource Estimation
    - Estimated man-hours by trade (Mechanical, Electrical, Instrumentation).
    
    Use a professional industrial tone with clear Markdown.
    """
    
    try:
        import openai
        client = openai.OpenAI(api_key=settings.openai_api_key)
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a world-class industrial maintenance consultant specializing in smelter operations."},
                {"role": "user", "content": prompt}
            ]
        )
        advice = response.choices[0].message.content
    except Exception as e:
        print(f"[OpenAI Error] {e}")
        # Fallback to Gemini if OpenAI fails
        from services.agent_manager import generate_with_retry
        advice = generate_with_retry(prompt)
    
    # 3. Save to Cache
    try:
        db_path = settings.db_path
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        cursor.execute("INSERT OR REPLACE INTO execution_advice_cache (wo_id, advice) VALUES (?, ?)", (wo_id, advice))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"Cache Error: {e}")

    return {"advice": advice, "asset_id": wo['asset_id'], "cached": False}

@app.post("/api/work-orders/{wo_id}/approve")
async def approve_work_order(wo_id: str):
    conn = sqlite3.connect(settings.db_path)
    cursor = conn.cursor()
    cursor.execute("UPDATE work_order SET work_order_status = 'In-Progress' WHERE id = ?", (wo_id,))
    
    # Also mark all associated permits as 'Available'
    cursor.execute("""
        UPDATE work_permit 
        SET status = 'Available' 
        WHERE work_order_task_item IN (
            SELECT id FROM work_order_task_item WHERE work_order = ?
        )
    """, (wo_id,))
    
    conn.commit()
    conn.close()
    return {"status": "success"}

@app.post("/api/work-permit/{permit_id}/generate")
async def generate_permit_document(permit_id: str):
    """Generate an AI-powered work permit PDF document using OpenAI."""
    import openai as _openai
    import json as _json

    _api_key = os.getenv("OPENAI_API_KEY", "")
    print(f"[Permit] Generating permit for: {permit_id} | API key present: {bool(_api_key)}")
    _openai_client = _openai.OpenAI(api_key=_api_key)

    # ── 1. Fetch permit details ──────────────────────────────────────────────
    permits = safe_query("""
        SELECT wp.id, wp.description, wp.type,
               wp.work_permit_open_day, wp.work_permit_open_time,
               wp.work_permit_end_day, wp.work_permit_end_time,
               wp.status, wp.status_change_timestamp,
               woti.work_order as work_order_id,
               woti.work_order_task_item_open_day as task_open_day,
               woti.work_order_task_item_open_time as task_open_time,
               woti.work_order_task_item_finish_day as task_finish_day,
               woti.work_order_task_item_finish_time as task_finish_time
        FROM work_permit wp
        JOIN work_order_task_item woti ON wp.work_order_task_item = woti.id
        WHERE wp.id = ?
    """, (permit_id,))

    if not permits:
        raise HTTPException(status_code=404, detail="Permit not found")
    permit = permits[0]

    # ── 2. Fetch related work order ──────────────────────────────────────────
    wo_id = permit.get("work_order_id", "")
    wos = safe_query("""
        SELECT w.id, w.repair_description, w.repair_type, w.work_order_class,
               w.work_order_status, w.work_order_open_day, w.work_order_open_time,
               a.name as asset_name, at.type as asset_type, a.location, a.criticality
        FROM work_order w
        LEFT JOIN work_order_task_item woti ON woti.work_order = w.id
        LEFT JOIN asset a ON woti.asset = a.id
        LEFT JOIN asset_type at ON a.type = at.id
        WHERE w.id = ?
        LIMIT 1
    """, (wo_id,))
    wo = wos[0] if wos else {}

    # ── 3. Fetch manpower ────────────────────────────────────────────────────
    manpower = safe_query("""
        SELECT te.name, te.role_designation, te.discipline_trade, tel.technician_service_period
        FROM technician_engineer_linkage tel
        JOIN work_order_task_item woti ON tel.work_order_task_item = woti.id
        JOIN technician_engineer te ON tel.technician_engineer_engaged = te.id
        WHERE woti.work_order = ?
    """, (wo_id,))

    # ── 4. Fetch materials ───────────────────────────────────────────────────
    materials = safe_query("""
        SELECT mm.description as material, m.quantity_used
        FROM task_material_linkage m
        JOIN work_order_task_item woti ON m.work_order_task_item = woti.id
        JOIN material_master mm ON m.material_used = mm.id
        WHERE woti.work_order = ?
    """, (wo_id,))

    # ── 5. Build LLM context ─────────────────────────────────────────────────
    personnel_lines = "\n".join([
        f"  - {m.get('name')} | {m.get('role_designation')} | {m.get('discipline_trade')} | {m.get('technician_service_period')}h"
        for m in manpower
    ])
    material_lines = "\n".join([
        f"  - {m.get('material')} x{m.get('quantity_used')}"
        for m in materials
    ])

    context = f"""
Work Permit ID: {permit.get('id')}
Permit Type: {permit.get('type')}
Description: {permit.get('description')}
Status: {permit.get('status')}
Valid From: {permit.get('work_permit_open_day')} {permit.get('work_permit_open_time')}
Valid Until: {permit.get('work_permit_end_day')} {permit.get('work_permit_end_time')}

Work Order: {wo_id}
Repair Type: {wo.get('repair_type')}
WO Class: {wo.get('work_order_class')}
Repair Description: {wo.get('repair_description')}
Asset: {wo.get('asset_name')} ({wo.get('asset_type')})
Location: {wo.get('location')}
Asset Criticality: {wo.get('criticality')}

Scheduled Window: {permit.get('task_open_day')} {permit.get('task_open_time')} to {permit.get('task_finish_day')} {permit.get('task_finish_time')}

Assigned Personnel ({len(manpower)}):
{personnel_lines}

Materials Required ({len(materials)}):
{material_lines}
"""

    # ── 6. Call OpenAI GPT-4o ─────────────────────────────────────────────────
    system_prompt = """You are a certified HSE (Health, Safety & Environment) officer and industrial maintenance expert 
specializing in aluminum smelter operations. Generate a comprehensive, professional work permit document.

Return ONLY valid JSON with exactly these keys:
{
  "permit_type_full": "Full official name of the permit type",
  "work_scope": "2-3 sentences describing the precise work scope",
  "location_details": "Specific location and access path description",
  "hazard_identification": ["4-6 specific hazards relevant to this work type and asset"],
  "risk_level": "HIGH or MEDIUM or LOW",
  "risk_justification": "One sentence justifying the risk level",
  "safety_controls": ["5-7 specific safety control measures for this work"],
  "ppe_requirements": ["All required PPE items with specifications"],
  "isolation_requirements": "Specific lockout/tagout and energy isolation procedure",
  "environmental_controls": ["2-3 environmental protection measures"],
  "emergency_procedure": "Step-by-step emergency response for this specific work type",
  "special_instructions": "Any specific technical or operational instructions",
  "authorization_conditions": "Conditions that must be met for permit validity",
  "competency_requirements": "Required certifications and qualifications for this work"
}"""

    try:
        print(f"[Permit] Calling OpenAI GPT-4o for permit {permit_id}...")
        response = _openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Generate a work permit document for:\n{context}"}
            ],
            response_format={"type": "json_object"}
        )
        ai_data = _json.loads(response.choices[0].message.content)
        print(f"[Permit] OpenAI response received successfully for {permit_id}")
    except Exception as e:
        print(f"[Permit Generation AI ERROR] {type(e).__name__}: {e} | USING MOCK FALLBACK")
        # ── Mock Fallback ─────────────────────────────────────────────────────
        ai_data = {
            "permit_type_full": f"OFFICIAL {permit.get('type', 'GENERAL')} PERMIT",
            "work_scope": f"Maintenance intervention for {wo.get('repair_description', 'specified equipment')}.",
            "location_details": f"Site location: {wo.get('location', 'Vedanta Jharsuguda')}",
            "hazard_identification": ["Mechanical movement/pinch points", "Potential for high-temperature contact", "Working at heights", "Trip and fall hazards"],
            "risk_level": "MEDIUM",
            "risk_justification": "Standard risk for routine maintenance activities on industrial equipment.",
            "safety_controls": ["Energy isolation (LOTO) verified", "Work area cordoned off", "Certified tools used", "Constant monitoring of activity"],
            "ppe_requirements": ["Hard Hat (ISI Mark)", "Safety Boots (S1P)", "High-Visibility Vest", "Mechanical Protection Gloves", "Safety Goggles"],
            "isolation_requirements": "Electrical isolation of drive motor and mechanical lockout of rotating components required.",
            "environmental_controls": ["Proper disposal of used lubricants", "Noise suppression if exceeding 85dB"],
            "emergency_procedure": "Stop work immediately, notify area manager, and evacuate via designated assembly point.",
            "special_instructions": "Ensure all tools are accounted for before finalizing permit.",
            "authorization_conditions": "Valid only for the specified shift and personnel.",
            "competency_requirements": "Trained maintenance technician with valid safety certification."
        }

    return {
        "permit": permit,
        "work_order": wo,
        "manpower": manpower,
        "materials": materials,
        "ai_document": ai_data
    }

@app.post("/api/work-permit/{permit_id}/download-docx")
async def download_permit_docx(permit_id: str, data: dict):
    """Convert provided permit JSON data into a Word document."""
    import io
    from docx import Document as DocxDocument
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ai_data = data.get("ai_document", {})
    permit = data.get("permit", {})
    manpower = data.get("manpower", [])

    doc = DocxDocument()
    
    # Title
    title = doc.add_heading("VEDANTA JHARSUGUDA — WORK PERMIT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph(f"{ai_data.get('permit_type_full', permit.get('type', ''))} PERMIT  |  {permit_id}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Issued: {permit.get('work_permit_open_day','')} {permit.get('work_permit_open_time','')}   |   Expires: {permit.get('work_permit_end_day','')} {permit.get('work_permit_end_time','')}   |   Risk: {ai_data.get('risk_level','')}")
    doc.add_paragraph()

    def add_section(heading, content):
        doc.add_heading(heading, level=2)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(f"• {item}", style="List Bullet")
        else:
            doc.add_paragraph(str(content))

    sections = [
        ("Work Scope", "work_scope"),
        ("Location Details", "location_details"),
        ("Risk Justification", "risk_justification"),
        ("Hazard Identification", "hazard_identification"),
        ("Safety Controls", "safety_controls"),
        ("PPE Requirements", "ppe_requirements"),
        ("Isolation / LOTO Requirements", "isolation_requirements"),
        ("Environmental Controls", "environmental_controls"),
        ("Emergency Procedure", "emergency_procedure"),
        ("Special Instructions", "special_instructions"),
        ("Authorization Conditions", "authorization_conditions"),
        ("Competency Requirements", "competency_requirements")
    ]

    for title_text, key in sections:
        add_section(title_text, ai_data.get(key, ""))

    # Personnel table
    doc.add_heading("Assigned Personnel", level=2)
    if manpower:
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Name", "Role", "Discipline", "Hours"
        for m in manpower:
            row = tbl.add_row().cells
            row[0].text = m.get("name","")
            row[1].text = m.get("role_designation","")
            row[2].text = m.get("discipline_trade","")
            row[3].text = str(m.get("technician_service_period",""))

    doc.add_paragraph()
    doc.add_heading("Signatures", level=2)
    for role in ["Permit Issuer", "Permit Receiver", "Safety Officer", "Area Manager"]:
        doc.add_paragraph(f"{role}: ___________________________    Date: ____________")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"WorkPermit_{permit_id}.docx\""}
    )

class ManpowerReasoningRequest(BaseModel):
    manpower: list

@app.post("/api/work-order/{wo_id}/manpower-reasoning")
async def get_manpower_reasoning(wo_id: str, payload: ManpowerReasoningRequest):
    log_stage(20, f"Generating manpower assignment reasoning for WO {wo_id}")
    
    # Fetch WO context
    wos = safe_query("""
        SELECT w.repair_description, w.repair_type, w.work_order_class,
               a.name as asset_name, at.type as asset_type
        FROM work_order w
        LEFT JOIN work_order_task_item woti ON woti.work_order = w.id
        LEFT JOIN asset a ON woti.asset = a.id
        LEFT JOIN asset_type at ON a.type = at.id
        WHERE w.id = ?
        LIMIT 1
    """, (wo_id,))
    
    wo = wos[0] if wos else {}
    manpower_list = payload.manpower

    if not manpower_list:
        return {}

    manpower_str = "\n".join([
        f"- ID: {m.get('technician_id', m.get('id', 'Unknown'))}, Name: {m.get('technician_name', m.get('name', 'Unknown'))}, Role: {m.get('role_designation', 'Unknown')}, Discipline: {m.get('discipline_trade', 'Unknown')}"
        for m in manpower_list
    ])

    prompt = f"""
    You are an expert maintenance planner at Vedanta Jharsuguda.
    Explain briefly (1-2 short sentences) WHY each of the following personnel was assigned to this work order.
    Make it sound highly analytical and intelligent (e.g. "Assigned due to their specialized electrical discipline matching the repair type").
    
    Work Order Context:
    Description: {wo.get('repair_description', 'N/A')}
    Type: {wo.get('repair_type', 'N/A')}
    Class: {wo.get('work_order_class', 'N/A')}
    Asset: {wo.get('asset_name', 'N/A')} ({wo.get('asset_type', 'N/A')})
    
    Assigned Personnel:
    {manpower_str}
    
    Return ONLY a valid JSON object mapping the exact ID to the reasoning string.
    Example: {{"T001": "Reasoning...", "E002": "Reasoning..."}}
    """
    
    from services.agent_manager import generate_with_retry
    import json
    
    try:
        response_text = generate_with_retry(prompt=prompt, system_prompt="You must output ONLY raw, valid JSON. Do not use markdown blocks like ```json.")
        # Attempt to parse json
        response_text = response_text.replace("```json", "").replace("```", "").strip()
        data = json.loads(response_text)
        return data
    except Exception as e:
        print(f"[Manpower Reasoning ERROR] {e}")
        # Return generic fallback
        return {m.get('technician_id', m.get('id', 'Unknown')): "Assigned to fulfill the standard operational and disciplinary requirements for this task." for m in manpower_list}

class MaterialReasoningRequest(BaseModel):
    materials: list

@app.post("/api/work-order/{wo_id}/material-reasoning")
async def get_material_reasoning(wo_id: str, payload: MaterialReasoningRequest):
    log_stage(21, f"Generating material selection reasoning for WO {wo_id}")
    
    # Fetch WO context
    wos = safe_query("""
        SELECT w.repair_description, w.repair_type, w.work_order_class,
               a.name as asset_name, at.type as asset_type
        FROM work_order w
        LEFT JOIN work_order_task_item woti ON woti.work_order = w.id
        LEFT JOIN asset a ON woti.asset = a.id
        LEFT JOIN asset_type at ON a.type = at.id
        WHERE w.id = ?
        LIMIT 1
    """, (wo_id,))
    
    wo = wos[0] if wos else {}
    materials_list = payload.materials

    if not materials_list:
        return {}

    materials_str = "\n".join([
        f"- Material: {m.get('material', 'Unknown')}, Qty: {m.get('recommended_quantity', 'Unknown')}"
        for m in materials_list
    ])

    prompt = f"""
    You are an expert maintenance planner at Vedanta Jharsuguda.
    Explain briefly (1 short sentence) WHY each of the following materials was selected for this work order.
    Make it sound highly analytical (e.g. "Required for the high-temperature bearing replacement").
    
    Work Order Context:
    Description: {wo.get('repair_description', 'N/A')}
    Type: {wo.get('repair_type', 'N/A')}
    Class: {wo.get('work_order_class', 'N/A')}
    Asset: {wo.get('asset_name', 'N/A')} ({wo.get('asset_type', 'N/A')})
    
    Selected Materials:
    {materials_str}
    
    Return ONLY a valid JSON object mapping the exact Material Name to the reasoning string.
    Example: {{"Bearing 6205": "Reasoning...", "Lubricant Grease": "Reasoning..."}}
    """
    
    from services.agent_manager import generate_with_retry
    import json
    
    try:
        response_text = generate_with_retry(prompt=prompt, system_prompt="You must output ONLY raw, valid JSON. Do not use markdown blocks like ```json.")
        # Attempt to parse json
        response_text = response_text.replace("```json", "").replace("```", "").strip()
        data = json.loads(response_text)
        return data
    except Exception as e:
        print(f"[Material Reasoning ERROR] {e}")
        # Return generic fallback
        return {m.get('material', 'Unknown'): "Required standard consumable for this type of repair task." for m in materials_list}

class ContractReasoningRequest(BaseModel):
    contracts: list

@app.post("/api/work-order/{wo_id}/contract-reasoning")
async def get_contract_reasoning(wo_id: str, payload: ContractReasoningRequest):
    log_stage(22, f"Generating contract selection reasoning for WO {wo_id}")
    
    # Fetch WO context
    wos = safe_query("""
        SELECT w.repair_description, w.repair_type, w.work_order_class,
               a.name as asset_name, at.type as asset_type
        FROM work_order w
        LEFT JOIN work_order_task_item woti ON woti.work_order = w.id
        LEFT JOIN asset a ON woti.asset = a.id
        LEFT JOIN asset_type at ON a.type = at.id
        WHERE w.id = ?
        LIMIT 1
    """, (wo_id,))
    
    wo = wos[0] if wos else {}
    contracts_list = payload.contracts

    if not contracts_list:
        return {}

    contracts_str = "\n".join([
        f"- Contract: {c.get('contract', 'Unknown')}, Type: {c.get('type', 'Unknown')}, Est. Value: {c.get('recommended_value', 'Unknown')}"
        for c in contracts_list
    ])

    prompt = f"""
    You are an expert maintenance planner at Vedanta Jharsuguda.
    Explain briefly (1 short sentence) WHY each of the following service contracts was selected for this work order.
    Make it sound highly analytical and cost-conscious (e.g. "Selected for specialized OEM support required during turbine overhauls").
    
    Work Order Context:
    Description: {wo.get('repair_description', 'N/A')}
    Type: {wo.get('repair_type', 'N/A')}
    Class: {wo.get('work_order_class', 'N/A')}
    Asset: {wo.get('asset_name', 'N/A')} ({wo.get('asset_type', 'N/A')})
    
    Selected Contracts:
    {contracts_str}
    
    Return ONLY a valid JSON object mapping the exact Contract Name to the reasoning string.
    Example: {{"OEM Maintenance Contract": "Reasoning...", "Heavy Lifting Service": "Reasoning..."}}
    """
    
    from services.agent_manager import generate_with_retry
    import json
    
    try:
        response_text = generate_with_retry(prompt=prompt, system_prompt="You must output ONLY raw, valid JSON. Do not use markdown blocks like ```json.")
        # Attempt to parse json
        response_text = response_text.replace("```json", "").replace("```", "").strip()
        data = json.loads(response_text)
        return data
    except Exception as e:
        print(f"[Contract Reasoning ERROR] {e}")
        # Return generic fallback
        return {c.get('contract', 'Unknown'): "Required for specialized external service support for this repair." for c in contracts_list}

class TaskReasoningRequest(BaseModel):
    tasks: list

@app.post("/api/work-order/{wo_id}/task-reasoning")
async def get_task_reasoning(wo_id: str, payload: TaskReasoningRequest):
    log_stage(23, f"Generating task duration reasoning for WO {wo_id}")
    
    # Fetch WO context
    wos = safe_query("""
        SELECT w.repair_description, w.repair_type, w.work_order_class,
               a.name as asset_name, at.type as asset_type
        FROM work_order w
        LEFT JOIN work_order_task_item woti ON woti.work_order = w.id
        LEFT JOIN asset a ON woti.asset = a.id
        LEFT JOIN asset_type at ON a.type = at.id
        WHERE w.id = ?
        LIMIT 1
    """, (wo_id,))
    
    wo = wos[0] if wos else {}
    tasks_list = payload.tasks

    if not tasks_list:
        return {}

    tasks_str = "\n".join([
        f"- Ref: {t.get('task_ref', 'Unknown')}, Task: {t.get('task_description', 'Unknown')}, Estimated Hours: {t.get('estimated_duration_hours', 8)}h"
        for t in tasks_list
    ])

    prompt = f"""
    You are an expert maintenance planner at Vedanta Jharsuguda.
    Explain briefly (1 short sentence) WHY each of the following tasks was assigned its specific estimated duration.
    Make it sound highly analytical and based on complexity/standard norms (e.g. "Estimated 12h due to complex alignment requirements and cooling period").
    
    Work Order Context:
    Description: {wo.get('repair_description', 'N/A')}
    Type: {wo.get('repair_type', 'N/A')}
    Class: {wo.get('work_order_class', 'N/A')}
    Asset: {wo.get('asset_name', 'N/A')} ({wo.get('asset_type', 'N/A')})
    
    Tasks & Estimated Durations:
    {tasks_str}
    
    Return ONLY a valid JSON object mapping the exact "Ref" value to the reasoning string.
    Example: {{"T001": "Reasoning...", "12345": "Reasoning..."}}
    """
    
    from services.agent_manager import generate_with_retry
    import json
    
    try:
        response_text = generate_with_retry(prompt=prompt, system_prompt="You must output ONLY raw, valid JSON. Do not use markdown blocks like ```json.")
        # Attempt to parse json
        response_text = response_text.replace("```json", "").replace("```", "").strip()
        data = json.loads(response_text)
        return data
    except Exception as e:
        print(f"[Task Reasoning ERROR] {e}")
        # Return generic fallback
        return {str(t.get('task_ref', 'Unknown')): "Duration based on standard man-hour norms for this discipline and task complexity." for t in tasks_list}

@app.get("/api/execution-plan/{work_order_id}")

async def get_execution_plan(work_order_id: str):
    tasks = safe_query("""
        SELECT wt.id, wt.task as task_ref, t.description as task_description, t.discipline,
               wt.work_order_task_item_open_day, wt.work_order_task_item_open_time,
               wt.work_order_task_item_finish_day, wt.work_order_task_item_finish_time,
               COALESCE(
                   -- Prefer technician/senior technician hours
                   (SELECT p.technician_service_period
                    FROM technician_engineer_linkage p
                    JOIN technician_engineer te ON te.id = p.technician_engineer_engaged
                    WHERE p.work_order_task_item = wt.id
                      AND LOWER(te.role_designation) IN ('technician', 'senior technician')
                    LIMIT 1),
                   -- Fallback to engineer hours
                   (SELECT p.technician_service_period
                    FROM technician_engineer_linkage p
                    JOIN technician_engineer te ON te.id = p.technician_engineer_engaged
                    WHERE p.work_order_task_item = wt.id
                    LIMIT 1),
                   -- Default 8h if no one assigned
                   8
               ) as estimated_duration_hours
        FROM work_order_task_item wt
        LEFT JOIN task t ON wt.task = t.id
        WHERE wt.work_order = ?
    """, (work_order_id,))
    
    materials = safe_query("""
        SELECT mm.description as material, m.quantity_used as recommended_quantity, m.material_price,
               m.lead_time, mm.id as material_id,
               COALESCE(SUM(inv.stock_available_on_hand), 0) as available_quantity
        FROM task_material_linkage m 
        JOIN work_order_task_item t ON m.work_order_task_item = t.id 
        JOIN material_master mm ON m.material_used = mm.id
        LEFT JOIN on_hand_inventory inv ON inv.material = mm.id
        WHERE t.work_order = ?
        GROUP BY mm.id, m.quantity_used, m.material_price, mm.description, m.lead_time
    """, (work_order_id,))

    # Estimate missing lead times if PR is needed
    for mat in materials:
        if mat['available_quantity'] < mat['recommended_quantity'] and mat['lead_time'] is None:
            log_stage(25, f"Estimating lead time for {mat['material']}")
            prompt = f"Estimate the typical industrial lead time in days for obtaining the following material for an aluminum smelter: '{mat['material']}'. Return ONLY a single integer representing the number of days."
            try:
                lt_str = generate_with_retry(prompt=prompt)
                # Extract digits
                lt = int(''.join(filter(str.isdigit, lt_str)))
                mat['lead_time'] = lt
                # Persist to DB
                conn = sqlite3.connect(settings.db_path)
                cursor = conn.cursor()
                cursor.execute("""
                    UPDATE task_material_linkage 
                    SET lead_time = ? 
                    WHERE material_used = ? AND work_order_task_item IN (
                        SELECT id FROM work_order_task_item WHERE work_order = ?
                    )
                """, (lt, mat['material_id'], work_order_id))
                conn.commit()
                conn.close()
            except:
                mat['lead_time'] = 7 # Default fallback
    
    manpower = safe_query("""
        SELECT p.technician_engineer_engaged as technician_id,
               te.name as technician_name,
               te.role_designation,
               te.discipline_trade,
               te.standard_hourly_rate,
               p.technician_service_period as service_period
        FROM technician_engineer_linkage p
        JOIN work_order_task_item t ON p.work_order_task_item = t.id
        LEFT JOIN technician_engineer te ON p.technician_engineer_engaged = te.id
        WHERE t.work_order = ?
    """, (work_order_id,))

    # ── Business Rule: Engineer must be paired with a Technician ──────────────
    engineer_roles = {"engineer", "senior engineer"}
    technician_roles = {"technician", "senior technician"}


    has_engineer = any(
        (m.get("role_designation") or "").lower() in engineer_roles
        for m in manpower
    )
    has_technician = any(
        (m.get("role_designation") or "").lower() in technician_roles
        for m in manpower
    )

    if has_engineer and not has_technician:
        # Find the discipline of the assigned engineer
        eng = next(
            m for m in manpower
            if (m.get("role_designation") or "").lower() in engineer_roles
        )
        discipline = eng.get("discipline_trade", "")

        # Pick the first available technician from the same discipline
        auto_tech = safe_query(
            """SELECT id, name, role_designation, discipline_trade, standard_hourly_rate
               FROM technician_engineer
               WHERE discipline_trade = ? AND role_designation IN ('Technician', 'Senior Technician')
               LIMIT 1""",
            (discipline,)
        )

        if auto_tech:
            tech = auto_tech[0]
            # Technician hours must be > engineer hours
            eng_hours = eng.get("service_period", 8)
            tech_hours = round(eng_hours * 1.5)
            manpower.append({
                "technician_id": tech["id"],
                "technician_name": tech["name"],
                "role_designation": tech["role_designation"],
                "discipline_trade": tech["discipline_trade"],
                "standard_hourly_rate": tech["standard_hourly_rate"],
                "service_period": tech_hours,
                "auto_assigned": True   # flag for UI labelling
            })

    # ── Business Rule: Engineer hours must be < Technician hours ─────────────
    for entry in manpower:
        role = (entry.get("role_designation") or "").lower()
        if role in engineer_roles:
            # Find paired technician(s) in same discipline
            same_disc_techs = [
                m for m in manpower
                if (m.get("role_designation") or "").lower() in technician_roles
                and m.get("discipline_trade") == entry.get("discipline_trade")
            ]
            if same_disc_techs:
                min_tech_hours = min(t["service_period"] for t in same_disc_techs)
                if entry["service_period"] >= min_tech_hours:
                    entry["service_period"] = max(1, round(min_tech_hours * 0.6))

    # ── Sync task estimated_duration_hours to final manpower hours ────────────
    # Task duration = technician hours (if any), else engineer hours, else 8h default
    tech_hours_final = None
    eng_hours_final = None
    for m in manpower:
        role = (m.get("role_designation") or "").lower()
        if role in technician_roles:
            # Use the highest technician hours (most representative of task effort)
            if tech_hours_final is None or m["service_period"] > tech_hours_final:
                tech_hours_final = m["service_period"]
        elif role in engineer_roles:
            if eng_hours_final is None or m["service_period"] > eng_hours_final:
                eng_hours_final = m["service_period"]

    resolved_duration = tech_hours_final if tech_hours_final is not None else (eng_hours_final if eng_hours_final is not None else 8)

    for task in tasks:
        task["estimated_duration_hours"] = resolved_duration

    contracts = safe_query("""
        SELECT c.contract_engaged as contract, 
               cs.contract_type as type,
               c.contract_value_expended as recommended_value,
               cs.contract_value as total_value
        FROM contract_linkage c 
        JOIN work_order_task_item t ON c.work_order_task_item = t.id 
        JOIN contract_services cs ON c.contract_engaged = cs.id
        WHERE t.work_order = ?
    """, (work_order_id,))
    
    work_permits = safe_query("""
        SELECT wp.id, wp.description, wp.type, wp.work_permit_open_day, wp.work_permit_open_time, 
               wp.work_permit_end_day, wp.work_permit_end_time, wp.status, wp.status_change_timestamp
        FROM work_permit wp
        JOIN work_order_task_item t ON wp.work_order_task_item = t.id
        WHERE t.work_order = ?
    """, (work_order_id,))

    # Enforce Rule: Status cannot be 'Unavailable' if it is issued (dates available)
    for wp in work_permits:
        if wp.get('work_permit_open_day') and wp.get('work_permit_end_day'):
            if wp.get('status') == 'Unavailable':
                wp['status'] = 'Available'
    
    # ── Live Auto-Closure Rule ────────────────────────────────────────────────
    # If the last task item's finish datetime < now and WO is still Pending/Approved,
    # close it immediately and stamp the closure time.
    _now = datetime.now()

    _last_task_finish = safe_query("""
        SELECT work_order_task_item_open_day   AS day,
               work_order_task_item_finish_time AS time
        FROM work_order_task_item
        WHERE work_order = ?
        ORDER BY work_order_task_item_open_day DESC,
                 work_order_task_item_finish_time DESC
        LIMIT 1
    """, (work_order_id,))

    if _last_task_finish:
        _lt = _last_task_finish[0]
        _day_str  = _lt.get("day", "") or ""
        _time_str = _lt.get("time", "") or ""
        try:
            _d, _m, _y = _day_str.strip().split('-')
            _h, _mi   = _time_str.strip().split(':')
            _last_dt  = datetime(2000 + int(_y), int(_m), int(_d), int(_h), int(_mi))
        except Exception:
            _last_dt = None

        if _last_dt and _last_dt < _now:
            # Only close if still in a non-final state
            _current_status = safe_query(
                "SELECT work_order_status FROM work_order WHERE id = ?", (work_order_id,)
            )
            _status_val = (_current_status[0].get("work_order_status", "") if _current_status else "").lower()
            if _status_val not in ("closed",):
                safe_query(
                    """UPDATE work_order
                       SET work_order_status   = 'Closed',
                           work_order_end_day  = ?,
                           work_order_end_time = ?
                       WHERE id = ?""",
                    (_day_str, _time_str, work_order_id)
                )
                # Ensure all permits for this WO are marked Available
                safe_query(
                    """UPDATE work_permit
                       SET status = 'Available'
                       WHERE work_order_task_item IN (
                           SELECT id FROM work_order_task_item WHERE work_order = ?
                       )""",
                    (work_order_id,)
                )

    work_order = safe_query("""
        SELECT w.*, a.name as asset_name, a.id as asset_id
        FROM work_order w
        LEFT JOIN work_order_task_item woti ON woti.work_order = w.id
        LEFT JOIN asset a ON woti.asset = a.id
        WHERE w.id = ?
        LIMIT 1
    """, (work_order_id,))
    wo_details = work_order[0] if work_order else {}
    
    total_manpower_cost = sum(m.get("service_period", 0) * m.get("standard_hourly_rate", 0) for m in manpower)
    total_material_cost = sum(m.get("recommended_quantity", 0) * m.get("material_price", 0) for m in materials)
    total_contract_cost = sum(c.get("recommended_value", 0) for c in contracts)
    total_cost = total_manpower_cost + total_material_cost + total_contract_cost
    
    estimated_cost = {
        "total": total_cost,
        "manpower": total_manpower_cost,
        "material": total_material_cost,
        "contract": total_contract_cost
    }
    
    return {
        "work_order": wo_details,
        "tasks": tasks,
        "materials": materials,
        "manpower": manpower,
        "contracts": contracts,
        "work_permits": work_permits,
        "estimated_cost": estimated_cost
    }

@app.post("/api/chat")
async def chat(request: ChatRequest):
    try:
        result = run_maintenance_assistant(request.message)
        # Log stages to terminal
        for log in result['stage_logs']:
            print(log)
        return result
    except Exception as e:
        print(f"[Chat Endpoint Error] {e}")
        return {
            "answer": f"Technical Error: {str(e)}",
            "data_used": {},
            "confidence": "low",
            "stage_logs": [f"Error: {str(e)}"]
        }

@app.post("/api/agent/{agent_id}")
async def run_agent(agent_id: str, request: ChatRequest):
    try:
        result = run_agent_workflow(agent_id, request.message)
        # Log stages to terminal
        for log in result['stage_logs']:
            print(log)
        return result
    except Exception as e:
        print(f"[Agent Endpoint Error] {e}")
        return {
            "answer": f"Technical Error: {str(e)}",
            "data_used": {},
            "confidence": "low",
            "stage_logs": [f"Error: {str(e)}"]
        }

# ─── WEBSOCKETS ───

@app.websocket("/ws/kpis")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()
    try:
        while True:
            kpis = get_all_kpis()
            await websocket.send_text(json.dumps([k.dict() if hasattr(k, 'dict') else k for k in kpis]))
            await asyncio.sleep(10)
    except WebSocketDisconnect:
        pass

@app.get("/api/schedule")
async def get_schedule(status: Optional[str] = None):
    now = datetime.now()

    # Build status filter for SQL
    status_filter = ""
    if status:
        status_filter = f"AND LOWER(w.work_order_status) = '{status.lower()}'"

    query = f"""
        SELECT
            w.id as work_order_id,
            w.repair_description as description,
            w.work_order_status,
            w.work_order_open_day as open_date,
            wt.id as task_id,
            wt.work_order_task_item_open_day as date,
            wt.work_order_task_item_open_time as start_time,
            wt.work_order_task_item_finish_time as end_time,
            te.id as technician_id,
            te.name as technician_name,
            te.role_designation,
            (SELECT COUNT(*) FROM work_permit wp
             JOIN work_order_task_item woti ON wp.work_order_task_item = woti.id
             WHERE woti.work_order = w.id) as permit_count
        FROM work_order w
        JOIN work_order_task_item wt ON w.id = wt.work_order
        LEFT JOIN technician_engineer_linkage tel ON wt.id = tel.work_order_task_item
        LEFT JOIN technician_engineer te ON tel.technician_engineer_engaged = te.id
        WHERE 1=1 {status_filter}
        ORDER BY wt.work_order_task_item_open_day, wt.work_order_task_item_open_time
    """
    results = safe_query(query)

    schedule = []
    seen_wo_ids = set()   # deduplicate — one calendar entry per WO

    for row in results:
        wo_id = row["work_order_id"]
        if wo_id in seen_wo_ids:
            continue

        db_date = row["date"] or ""
        if not db_date:
            continue

        try:
            parts = db_date.split('-')
            dt_obj = datetime(2000 + int(parts[2]), int(parts[1]), int(parts[0]))
        except Exception:
            continue

        seen_wo_ids.add(wo_id)
        item_status = row["work_order_status"] or "Pending"
        end_time_str = row["end_time"] or "23:59"
            
        schedule.append({
            "id": row["work_order_id"],
            "title": row["description"],
            "date": dt_obj.strftime("%Y-%m-%d"),
            "start": row["start_time"] or "00:00",
            "end": end_time_str,
            "technician": row["technician_name"],
            "technicianId": row["technician_id"],
            "role": row["role_designation"],
            "hasPermit": row["permit_count"] > 0,
            "status": item_status,
            "openDate": row["open_date"]
        })
    return schedule

@app.get("/api/system-updates")
async def get_system_updates():
    """Return a list of recent platform improvements and new features."""
    return [
        {
            "id": "word-permits",
            "title": "Interactive Word Permits",
            "description": "Permits now open in a premium preview pop-up with a dedicated 'Download as Word' export option.",
            "category": "New Feature",
            "timestamp": "Today"
        },
        {
            "id": "pending-rule",
            "title": "20% Pending Volume Floor",
            "description": "System now automatically maintains at least 20% work order volume in 'Pending' status to ensure a healthy maintenance pipeline.",
            "category": "System Rule",
            "timestamp": "Today"
        },
        {
            "id": "scheduling-integrity",
            "title": "Dual-Date Scheduling",
            "description": "Work order open dates are now backdated (Today-2/3) while task items are scheduled forward with a 4 WO/day capacity cap.",
            "category": "Integrity",
            "timestamp": "Yesterday"
        }
    ]

@app.post("/api/material-reservation/generate")
async def generate_material_reservation(data: dict):
    """Generate an AI-powered Material Reservation (MR) document."""
    import openai as _openai
    import json as _json

    _api_key = os.getenv("OPENAI_API_KEY", "")
    _openai_client = _openai.OpenAI(api_key=_api_key)

    material = data.get("material", "Unknown Material")
    qty = data.get("quantity", 0)
    wo_id = data.get("work_order_id", "N/A")
    asset_id = data.get("asset_id", "N/A")
    asset_name = data.get("asset_name", "N/A")

    # 1. Calculate Reservation Period (Start: First task start, End: Start + 2 days)
    _res = safe_query("""
        SELECT work_order_task_item_open_day as day,
               work_order_task_item_open_time as time
        FROM work_order_task_item
        WHERE work_order = ?
        ORDER BY work_order_task_item_open_day ASC,
                 work_order_task_item_open_time ASC
        LIMIT 1
    """, (wo_id,))
    
    start_date_str = "N/A"
    end_date_str = "N/A"
    if _res:
        _d_str = _res[0].get("day", "")
        _t_str = _res[0].get("time", "08:00")
        try:
            # Parse DD-MM-YY
            _d, _m, _y = _d_str.strip().split('-')
            start_dt = datetime(2000 + int(_y), int(_m), int(_d))
            end_dt = start_dt + timedelta(days=2)
            start_date_str = f"{start_dt.strftime('%d-%m-%y')} {_t_str}"
            end_date_str = f"{end_dt.strftime('%d-%m-%y')} 17:00"
        except Exception as e:
            print(f"Error parsing date for MR: {e}")

    system_prompt = f"""You are a materials manager at the Vedanta Jharsuguda Aluminum Plant.
    Generate a professional Material Reservation (MR) document.
    
    Return ONLY valid JSON with these keys:
    {{
      "mr_number": "MR-{wo_id[-4:]}-2026",
      "reservation_type": "Maintenance/Emergency",
      "material_specifications": "Detailed technical specs for {material} in the context of {asset_name} ({asset_id}) at Vedanta Jharsuguda",
      "storage_conditions": "Specific storage instructions for a tropical smelter environment",
      "handling_instructions": "Safety and handling for this specific item in a smelter",
      "criticality_impact": "Operational impact if this material is delayed for Work Order {wo_id} on {asset_id}",
      "warehouse_instructions": "Specific instructions for Vedanta's warehouse personnel",
      "validity_period": "{start_date_str} to {end_date_str}",
      "terms": "Standard Vedanta Jharsuguda material handling terms"
    }}"""

    prompt = f"Create a Material Reservation for {qty} units of '{material}' for Work Order {wo_id} (Asset: {asset_id} - {asset_name}). Period: {start_date_str} to {end_date_str}."

    try:
        response = _openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"}
        )
        ai_data = _json.loads(response.choices[0].message.content)
        return ai_data
    except Exception as e:
        print(f"[MR Generation AI ERROR] {e} | USING MOCK FALLBACK")
        return {
            "mr_number": f"MR-{wo_id[-4:]}-2026",
            "reservation_type": "Maintenance/Standard",
            "material_specifications": f"Technical specifications for {material} matching {asset_name} ({asset_id}) OEM requirements.",
            "storage_conditions": "Store in a cool, dry place away from direct heat and chemical exposure.",
            "handling_instructions": "Use appropriate lifting equipment for heavy components. Wear protective gloves.",
            "criticality_impact": f"High risk of operational downtime for {asset_name} if material delivery is delayed.",
            "warehouse_instructions": "Release to maintenance team lead after verification of Work Order {wo_id}.",
            "validity_period": f"{start_date_str} to {end_date_str}",
            "terms": "Subject to standard Vedanta Jharsuguda material handling and audit terms."
        }

@app.post("/api/purchase-requisition/generate")
async def generate_purchase_requisition(data: dict):
    """Generate an AI-powered Purchase Requisition (PR) document."""
    import openai as _openai
    import json as _json

    _api_key = os.getenv("OPENAI_API_KEY", "")
    _openai_client = _openai.OpenAI(api_key=_api_key)

    material = data.get("material", "Unknown Material")
    qty = data.get("quantity", 0)
    wo_id = data.get("work_order_id", "N/A")
    asset_id = data.get("asset_id", "N/A")
    asset_name = data.get("asset_name", "N/A")

    system_prompt = f"""You are a Senior Maintenance Engineer and Procurement Specialist at the Vedanta Jharsuguda Aluminum Smelter.
    Generate a high-detail, technically rigorous Purchase Requisition (PR) for internal engineering approval.
    
    The 'justification' field MUST be extremely detailed (at least 150 words). It should:
    1. Analyze the specific failure mode or degradation pattern of the '{material}' within the '{asset_name}' ({asset_id}).
    2. Cite technical consequences of delay, such as production loss in the Potline, safety risks in the Cast House, or environmental non-compliance in the FTP (Fume Treatment Plant).
    3. Explain the technical necessity of the requested specifications over standard alternatives.
    
    Use industry-specific terminology related to aluminum smelting (e.g., cryolite corrosion, magnetic field interference, thermal cycling, or alumina handling).
    
    Return ONLY valid JSON with these keys:
    {{
      "pr_number": "PR-{wo_id[-4:]}-2026",
      "requester_department": "Maintenance Division — Smelter Operations",
      "justification": "Detailed 2-3 paragraph technical justification",
      "technical_specifications": "High-precision technical standards and tolerances for {material}",
      "vendor_recommendations": ["List 2-3 globally recognized OEMs or certified local suppliers"],
      "estimated_budget": "Estimated cost for {qty} units in INR (₹)",
      "delivery_urgency": "IMMEDIATE (Production Critical) or STANDARD",
      "inspection_requirements": "Detailed QA/QC checklist and material test certificates (MTC) required",
      "approval_workflow": "Engineering Manager -> Operations Head -> Finance Director"
    }}"""

    prompt = f"Create a Purchase Requisition for {qty} units of '{material}' for Work Order {wo_id} (Asset: {asset_id} - {asset_name})."

    try:
        response = _openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"}
        )
        ai_data = _json.loads(response.choices[0].message.content)
        return ai_data
    except Exception as e:
        print(f"[PR Generation AI ERROR] {e} | USING MOCK FALLBACK")
        return {
            "pr_number": f"PR-{wo_id[-4:]}-2026",
            "requester_department": "Maintenance Division",

            "justification": f"Urgent replacement of {material} to restore {asset_name} operational integrity.",
            "technical_specifications": "Standard industrial grade specifications for aluminum smelter environments.",
            "vendor_recommendations": ["Global Industrial Supplies", "Vedanta Approved Local Vendors"],
            "estimated_budget": "₹85,000.00 (Estimated)",
            "delivery_urgency": "IMMEDIATE",
            "inspection_requirements": "Visual inspection for transit damage and verification of mill test certificates.",
            "approval_workflow": "Maintenance Lead -> Procurement Manager -> Finance Head"
        }

@app.post("/api/purchase-requisition/download-docx")
async def download_pr_docx(data: dict):
    """Convert provided PR JSON data into a Word document."""
    import io
    from docx import Document as DocxDocument
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pr_no = data.get("pr_number", "PR-NEW")
    mat = data.get("matName", "Unknown Material")

    doc = DocxDocument()
    
    # Title
    title = doc.add_heading("VEDANTA JHARSUGUDA — PURCHASE REQUISITION", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph(f"PR NUMBER: {pr_no}  |  MATERIAL: {mat}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Department: {data.get('requester_department','')}   |   Urgency: {data.get('delivery_urgency','')}   |   Budget: {data.get('estimated_budget','')}")
    doc.add_paragraph()

    def add_section(heading, content):
        doc.add_heading(heading, level=2)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(f"• {item}", style="List Bullet")
        else:
            doc.add_paragraph(str(content))

    sections = [
        ("Technical Justification", "justification"),
        ("Technical Specifications", "technical_specifications"),
        ("Vendor Recommendations", "vendor_recommendations"),
        ("QA/QC & Inspection Requirements", "inspection_requirements"),
        ("Approval Workflow", "approval_workflow")
    ]

    for title_text, key in sections:
        add_section(title_text, data.get(key, ""))

    doc.add_paragraph()
    doc.add_heading("Authorization Signatures", level=2)
    for role in ["Maintenance Lead", "Procurement Manager", "Finance Director", "Smelter Operations Head"]:
        doc.add_paragraph(f"{role}: ___________________________    Date: ____________")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"PurchaseRequisition_{pr_no}.docx\""}
    )

class UpdateLeadTimeRequest(BaseModel):
    material_id: str
    lead_time: int

@app.post("/api/work-order/{wo_id}/update-material-lead-time")
async def update_material_lead_time(wo_id: str, payload: UpdateLeadTimeRequest):
    conn = sqlite3.connect(settings.db_path)
    cursor = conn.cursor()
    cursor.execute("""
        UPDATE task_material_linkage 
        SET lead_time = ? 
        WHERE material_used = ? AND work_order_task_item IN (
            SELECT id FROM work_order_task_item WHERE work_order = ?
        )
    """, (payload.lead_time, payload.material_id, wo_id))
    conn.commit()
    conn.close()
    
    # Reschedule the work order
    new_date = reschedule_work_order(wo_id)
    return {"status": "success", "new_date": new_date}

@app.get("/api/inventory/summary")
async def get_inventory_summary():
    conn = sqlite3.connect(settings.db_path)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # 1. Total Value in INR
    cursor.execute("""
        SELECT SUM(oh.stock_available_on_hand * mp.price_per_unit) as total_value
        FROM on_hand_inventory oh
        JOIN material_price mp ON oh.material = mp.material
    """)
    total_value = cursor.fetchone()['total_value'] or 0
    
    # 2. Pending PR Count
    cursor.execute("SELECT COUNT(*) as count FROM purchase_requisition WHERE LOWER(status) = 'pending'")
    pending_pr_count = cursor.fetchone()['count'] or 0
    
    # 3. Critical Spares Out of Stock (Spares with 0 stock)
    cursor.execute("""
        SELECT COUNT(DISTINCT mm.id) as count
        FROM material_master mm
        JOIN on_hand_inventory oh ON mm.id = oh.material
        WHERE mm.material_type = 'Spares' AND oh.stock_available_on_hand <= 0
    """)
    critical_oos_count = cursor.fetchone()['count'] or 0
    
    # 4. Obsolescence Count
    cursor.execute("""
        SELECT oh.receipt_date, mm.shelf_life
        FROM on_hand_inventory oh
        JOIN material_master mm ON oh.material = mm.id
        WHERE mm.shelf_life IS NOT NULL AND oh.receipt_date IS NOT NULL
    """)
    rows = cursor.fetchall()
    obsolescence_count = 0
    from datetime import datetime, timedelta
    today = datetime.now()
    threshold = today + timedelta(days=30)
    
    for row in rows:
        try:
            r_date = datetime.strptime(row['receipt_date'], "%d-%m-%y")
            expiry_date = r_date + timedelta(days=row['shelf_life'])
            if expiry_date <= threshold:
                obsolescence_count += 1
        except:
            continue

    conn.close()
    return {
        "total_value_inr": total_value,
        "pending_pr_count": pending_pr_count,
        "critical_oos_count": critical_oos_count,
        "obsolescence_count": obsolescence_count
    }

@app.get("/api/inventory/pending-prs")
async def get_pending_prs():
    conn = sqlite3.connect(settings.db_path)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("""
        SELECT 
            pr.id, 
            mm.description as material_name, 
            pr.status,
            mp.price_per_unit as unit_price
        FROM purchase_requisition pr
        JOIN material_master mm ON pr.material = mm.id
        LEFT JOIN material_price mp ON mm.id = mp.id
        WHERE LOWER(pr.status) = 'pending'
    """)
    prs = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return prs

@app.get("/api/inventory/critical-spares-oos")
async def get_critical_spares_oos():
    conn = sqlite3.connect(settings.db_path)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # Find spares with 0 stock, related asset, and PR status
    cursor.execute("""
        SELECT 
            mm.id as material_id,
            mm.description as material_name,
            oh.stock_available_on_hand as stock,
            a.name as asset_name,
            a.id as asset_id,
            (SELECT status FROM purchase_requisition WHERE material = mm.id LIMIT 1) as pr_status
        FROM material_master mm
        JOIN on_hand_inventory oh ON mm.id = oh.material
        LEFT JOIN task_material_linkage tml ON mm.id = tml.material_used
        LEFT JOIN work_order_task_item woti ON tml.work_order_task_item = woti.id
        LEFT JOIN asset a ON woti.asset = a.id
        WHERE mm.material_type = 'Spares' AND oh.stock_available_on_hand <= 0
        GROUP BY mm.id
    """)
    items = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return items

@app.get("/api/inventory/obsolescence")
async def get_obsolescence_data():
    conn = sqlite3.connect(settings.db_path)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT 
            oh.material as material_id,
            mm.description as material_name,
            oh.receipt_date,
            mm.shelf_life,
            oh.stock_available_on_hand as stock
        FROM on_hand_inventory oh
        JOIN material_master mm ON oh.material = mm.id
        WHERE mm.shelf_life IS NOT NULL AND oh.receipt_date IS NOT NULL
    """)
    rows = cursor.fetchall()
    results = []
    from datetime import datetime, timedelta
    today = datetime.now()
    threshold = today + timedelta(days=30)
    
    for row in rows:
        try:
            r_date = datetime.strptime(row['receipt_date'], "%d-%m-%y")
            expiry_date = r_date + timedelta(days=row['shelf_life'])
            if expiry_date <= threshold:
                item = dict(row)
                item['expiry_date'] = expiry_date.strftime("%d-%m-%y")
                item['days_remaining'] = (expiry_date - today).days
                results.append(item)
        except:
            continue
            
    conn.close()
    return results

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
